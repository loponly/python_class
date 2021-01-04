from docx import Document
from docx.shared import Inches


doc = Document()

doc.add_heading('Document', 0)

p = doc.add_paragraph('Some text')
p.add_run(' bold text ').bold = True
p.add_run(' italic text ').italic = True


doc.save('sample.docx')
